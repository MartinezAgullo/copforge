"""
Document Processing Tools - PDF, TXT, DOCX extraction.
"""

import logging
import os
from datetime import datetime
from pathlib import Path
from typing import Any

from src.core.telemetry import get_tracer, traced_operation

logger = logging.getLogger(__name__)
tracer = get_tracer("copforge.mcp.multimodal.document")

SUPPORTED_DOCUMENT_EXTENSIONS = {".txt", ".pdf", ".docx"}


def is_document_file(file_path: str) -> bool:
    """Check if file is a supported document format."""
    return Path(file_path).suffix.lower() in SUPPORTED_DOCUMENT_EXTENSIONS


def validate_document_file(document_path: str) -> tuple[bool, str | None]:
    """Validate document file exists and is supported."""
    if not os.path.exists(document_path):
        return False, f"Document file not found: {document_path}"
    extension = Path(document_path).suffix.lower()
    if extension not in SUPPORTED_DOCUMENT_EXTENSIONS:
        return False, f"Unsupported format: {extension}. Supported: {SUPPORTED_DOCUMENT_EXTENSIONS}"
    if extension == ".doc":
        return False, "Legacy .doc format not supported. Convert to .docx or .pdf"
    file_size_mb = os.path.getsize(document_path) / (1024 * 1024)
    if file_size_mb > 50:
        return False, f"Document too large: {file_size_mb:.1f}MB (max 50MB)"
    return True, None


def _extract_text_from_pdf(pdf_path: str) -> str | None:
    """Extract text from PDF using PyPDF2."""
    try:
        import PyPDF2

        text_content = []
        with open(pdf_path, "rb") as file:
            reader = PyPDF2.PdfReader(file)
            for page_num, page in enumerate(reader.pages):
                text = page.extract_text()
                if text.strip():
                    text_content.append(f"--- Page {page_num + 1} ---\n{text}")
        return "\n\n".join(text_content)
    except ImportError as e:
        raise ImportError("PyPDF2 not installed. Install with: pip install PyPDF2") from e
    except Exception as e:
        logger.error(f"PDF extraction error: {e}")
        return None


def _read_text_file(text_path: str) -> str | None:
    """Read plain text file with fallback encodings."""
    encodings = ["utf-8", "latin-1", "cp1252", "iso-8859-1"]
    for enc in encodings:
        try:
            with open(text_path, encoding=enc) as file:
                return file.read()
        except (UnicodeDecodeError, LookupError):
            continue
    logger.error(f"Failed to read {text_path} with any encoding")
    return None


def _extract_text_from_docx(docx_path: str) -> str | None:
    """Extract text from DOCX using python-docx."""
    try:
        import docx

        doc = docx.Document(docx_path)
        paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]

        # Extract tables
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    tables_text.append(row_text)

        full_text = "\n\n".join(paragraphs)
        if tables_text:
            full_text += "\n\n--- TABLES ---\n" + "\n".join(tables_text)
        return full_text
    except ImportError as e:
        raise ImportError("python-docx not installed. Install with: pip install python-docx") from e
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        return None


def clean_extracted_text(text: str, max_lines: int | None = None) -> str:
    """Clean and preprocess extracted text."""
    if not text:
        return ""

    lines = text.split("\n")
    cleaned_lines: list[str] = []

    for line in lines:
        line = line.strip()
        if not line:
            if cleaned_lines and cleaned_lines[-1] != "":
                cleaned_lines.append("")
            continue
        line = " ".join(line.split())
        cleaned_lines.append(line)

    if max_lines and len(cleaned_lines) > max_lines:
        cleaned_lines = cleaned_lines[:max_lines]
        cleaned_lines.append(f"\n... [Truncated - total {len(lines)} lines] ...")

    return "\n".join(cleaned_lines)


def extract_text_from_document(
    document_path: str,
    clean_text: bool = True,
    max_lines: int | None = None,
) -> dict[str, Any]:
    """
    Extract text from document (auto-detects format).

    Args:
        document_path: Path to document file
        clean_text: Whether to clean/preprocess text
        max_lines: Maximum lines to extract (None = all)

    Returns:
        Dict with extraction result
    """
    with traced_operation(
        tracer, "extract_text_from_document", {"document_path": document_path}
    ) as span:
        try:
            is_valid, error = validate_document_file(document_path)
            if not is_valid:
                return {
                    "success": False,
                    "text": "",
                    "format": None,
                    "num_pages": None,
                    "num_lines": 0,
                    "error": error,
                }

            extension = Path(document_path).suffix.lower()
            logger.info(f"Extracting text from: {Path(document_path).name}")

            if extension == ".pdf":
                raw_text = _extract_text_from_pdf(document_path)
                doc_format = "pdf"
            elif extension == ".txt":
                raw_text = _read_text_file(document_path)
                doc_format = "txt"
            elif extension == ".docx":
                raw_text = _extract_text_from_docx(document_path)
                doc_format = "docx"
            else:
                return {
                    "success": False,
                    "text": "",
                    "format": None,
                    "num_pages": None,
                    "num_lines": 0,
                    "error": f"Unsupported format: {extension}",
                }

            if raw_text is None:
                return {
                    "success": False,
                    "text": "",
                    "format": doc_format,
                    "num_pages": None,
                    "num_lines": 0,
                    "error": "Failed to extract text from document",
                }

            final_text = clean_extracted_text(raw_text, max_lines) if clean_text else raw_text
            num_lines = len(final_text.split("\n"))

            span.set_attribute("document.format", doc_format)
            span.set_attribute("document.lines", num_lines)

            return {
                "success": True,
                "text": final_text,
                "format": doc_format,
                "num_pages": None,  # Could be enhanced for PDFs
                "num_lines": num_lines,
                "error": None,
            }
        except Exception as e:
            logger.exception("Document extraction failed")
            return {
                "success": False,
                "text": "",
                "format": None,
                "num_pages": None,
                "num_lines": 0,
                "error": f"Extraction failed: {e!s}",
            }


def process_document(document_path: str, max_lines: int | None = 1000) -> dict[str, Any]:
    """
    High-level document processing function.

    Args:
        document_path: Path to document file
        max_lines: Maximum lines to extract

    Returns:
        Dict with document extraction result
    """
    with traced_operation(tracer, "process_document") as span:
        file_name = Path(document_path).name
        result = extract_text_from_document(
            document_path=document_path, clean_text=True, max_lines=max_lines
        )

        span.set_attribute("document.success", result["success"])

        if not result["success"]:
            return {
                "success": False,
                "file_name": file_name,
                "format": None,
                "num_lines": 0,
                "text": "",
                "error": result["error"],
            }

        if not result["text"].strip():
            return {
                "success": False,
                "file_name": file_name,
                "format": result["format"],
                "num_lines": 0,
                "text": "",
                "error": "Document appears empty or contains no extractable text",
            }

        return {
            "success": True,
            "file_name": file_name,
            "format": result["format"],
            "num_lines": result["num_lines"],
            "text": result["text"],
            "error": None,
        }


def get_document_info(document_path: str) -> dict[str, Any] | None:
    """Get basic document metadata without extracting content."""
    try:
        if not os.path.exists(document_path):
            return None
        file_path = Path(document_path)
        file_size_mb = os.path.getsize(document_path) / (1024 * 1024)
        info: dict[str, Any] = {
            "file_name": file_path.name,
            "file_extension": file_path.suffix.lower(),
            "file_size_mb": round(file_size_mb, 2),
            "file_modified": datetime.fromtimestamp(os.path.getmtime(document_path)).isoformat(),
        }
        if file_path.suffix.lower() == ".pdf":
            try:
                import PyPDF2

                with open(document_path, "rb") as file:
                    reader = PyPDF2.PdfReader(file)
                    info["num_pages"] = len(reader.pages)
            except Exception:
                info["num_pages"] = None
        return info
    except Exception:
        return None
